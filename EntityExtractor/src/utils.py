#!/usr/bin/python
# -*- coding: utf-8 -*-

"""
Created by Lalitha Madhuri Putchala on Dec 10 2017
"""

import nltk
from nltk import word_tokenize, pos_tag, ne_chunk
from nltk.tokenize.punkt import PunktSentenceTokenizer
from docx.enum.text import WD_COLOR_INDEX


def highlight_text(doc, word):
    ''' 
        this function obtains the content of each paragraph and checks for the given word
        if the given word matches it highlights with yellow
    '''

    for paragraph in doc.paragraphs:
        if word in paragraph.text:
            substrings = paragraph.text.split(word)
            for substring in substrings[:-1]:
                paragraph.text = substring
                font = paragraph.add_run(word).font
                font.highlight_color = WD_COLOR_INDEX.YELLOW
            paragraph.add_run(substrings[-1])


def get_sentence_tokens(text):
    paragraphs = [p for p in text.split('\n') if p]
    tokenizer = PunktSentenceTokenizer()
    sentence_tokens = [tokenizer.tokenize(paragraph) for paragraph in
                       paragraphs]
    return sentence_tokens


def get_word_tokens(text):
    return word_tokenize(text)


def get_pos_tags(text):
    return nltk.pos_tag(text)


def get_ne_chunks(pos):
    return ne_chunk(pos)

